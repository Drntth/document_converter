"""
Feldolgozási logika: dokumentum előfeldolgozás, konverzió, gazdagítás, export.
Minden lépés külön függvényben, abszolút importokkal, PEP 257 docstringekkel.
"""

import json
import os
import time
from concurrent.futures import ThreadPoolExecutor
from functools import partial
from pathlib import Path
from typing import Dict

import config.settings
from config.logging_config import structlog_logger
from docling.document_converter import DocumentConverter
from docx import Document
from processing.enrichment import summarize_image, summarize_table
from processing.postprocessor import (
    extract_abbreviations,
    insert_abbreviations,
    insert_footnotes,
    prepend_abbreviation_section,
    remove_abbreviation_phrases,
    remove_footnote_references,
)
from processing.preprocessor import (
    remove_empty_pages,
    remove_toc_by_field,
    remove_toc_by_paragraphs,
    remove_toc_by_table,
    remove_toc_by_text,
    remove_toc_by_xml,
)
from services.file_service import save_abbreviations, validate_document, write_json
from unstructured.partition.auto import partition


def preprocess_docx(
    source_file: Path,
    remove_headers: bool,
    remove_footers: bool,
    remove_toc: bool,
    remove_empty: bool,
) -> Path:
    """
    Fő előfeldolgozó függvény választható lépésekkel.

    Args:
        source_file (Path): A bemeneti dokumentum elérési útja.
        remove_headers (bool): Fejlécek eltávolítása.
        remove_footers (bool): Láblécek eltávolítása.
        remove_toc (bool): Tartalomjegyzék eltávolítása.
        remove_empty (bool): Üres oldalak/sorok eltávolítása.

    Returns:
        Path: Az előfeldolgozott dokumentum elérési útja.
    """

    logger = structlog_logger.bind(source_file=source_file)
    start_time = time.time()
    logger.info(
        "DOCX előfeldolgozás elkezdődött", context={"source_file": str(source_file)}
    )

    headers_removed = False
    footers_removed = False
    toc_removed = False
    empty_removed = 0

    doc = Document(source_file)

    if remove_headers:
        for section in doc.sections:
            for header in section.header.paragraphs:
                header.clear()
        headers_removed = True
        logger.debug("Fejlécek eltávolítva")

    if remove_footers:
        for section in doc.sections:
            for footer in section.footer.paragraphs:
                footer.clear()
        footers_removed = True
        logger.debug("Láblécek eltávolítva")

    if remove_toc:
        methods = [
            (remove_toc_by_xml, "XML alapú"),
            (remove_toc_by_field, "Mezőkód alapú"),
            (remove_toc_by_paragraphs, "Bekezdés alapú"),
            (
                partial(remove_toc_by_text, keyword="tartalomjegyzék"),
                "Szöveges keresés",
            ),
            (remove_toc_by_table, "Táblázat alapú"),
        ]

        for method, method_name in methods:
            try:
                if method(doc, logger):
                    logger.info(f"Tartalomjegyzék eltávolítva: {method_name} módszer")
                    toc_removed = True
                    # break
                else:
                    logger.debug(
                        f"Nem található tartalomjegyzék: {method_name} módszer"
                    )
            except Exception as e:
                logger.error(
                    f"Hiba a {method_name} módszer során: {str(e)}", exc_info=True
                )

        logger.debug("Eredeti TOC eltávolítási logika indítása")

        toc_heading = None
        toc_start = None
        toc_end = None

        for i, para in enumerate(doc.paragraphs):
            if "tartalomjegyzék" in para.text.lower():
                toc_heading = para
                toc_start = i
                logger.debug(f"Tartalomjegyzék címsor megtalálva a {i}. bekezdésben")
                break

        if toc_start is not None:
            for i in range(toc_start + 1, len(doc.paragraphs)):
                para = doc.paragraphs[i]
                if not para.text.strip() or any(
                    run.text == "\x0c" for run in para.runs
                ):
                    toc_end = i
                    logger.debug(f"Tartalomjegyzék vége megtalálva a {i}. bekezdésben")
                    break
                elif para.style.name.startswith("Heading"):
                    toc_end = i
                    logger.debug(f"Tartalomjegyzék vége új címsornál (bekezdés {i})")
                    break

            if toc_end is not None:
                for i in range(toc_start, toc_end + 1):
                    if toc_start < len(doc.paragraphs):
                        doc.paragraphs[toc_start]._element.getparent().remove(
                            doc.paragraphs[toc_start]._element
                        )
                toc_removed = True
                logger.debug("Tartalomjegyzék eltávolítva (eredeti logika)")
            else:
                toc_heading._element.getparent().remove(toc_heading._element)
                toc_removed = True
                logger.debug("Tartalomjegyzék címsor eltávolítva (vége nem található)")

    empty_removed = remove_empty_pages(doc, logger, remove_empty)

    doc.save(source_file)

    logger.info(
        "DOCX előfeldolgozás kész",
        context={
            "source_file": str(source_file),
            "headers_removed": headers_removed,
            "footers_removed": footers_removed,
            "toc_removed": toc_removed,
            "empty_removed": empty_removed,
            "duration": time.time() - start_time,
        },
    )

    return source_file


def postprocess_input_file(
    source_file: Path, abbreviation_strategy: str, footnote_handling: str
) -> Path:
    """
    Dokumentum utófeldolgozása a preprocess után.

    A művelet tartalmazza a rövidítések és lábjegyzetek kezelését.

    Args:
        source_file (Path): A dokumentum elérési útja.
        abbreviation_strategy (str): Rövidítések kezelési stratégia.
        footnote_handling (str): Lábjegyzetek kezelési stratégia.

    Returns:
        Path: A feldolgozott dokumentum elérési útja.
    """
    logger = structlog_logger.bind(source_file=str(source_file))
    logger.info("Utófeldolgozás indítása", context={"source_file": str(source_file)})

    paths = config.settings.get_paths_for_file(str(source_file))

    try:
        doc = Document(source_file)

        if abbreviation_strategy != "none":
            abbreviations = extract_abbreviations(doc)
            save_abbreviations(abbreviations, paths["abbreviations_output"])

            if abbreviations:
                remove_abbreviation_phrases(doc, abbreviations)

                logger.info(
                    f"{len(abbreviations)} darab feldolgozandó rövidítés található"
                )

                if abbreviation_strategy == "inline":
                    insert_abbreviations(doc, abbreviations)
                elif abbreviation_strategy == "section":
                    prepend_abbreviation_section(doc, abbreviations)

        if footnote_handling == "remove":
            remove_footnote_references(doc)
            logger.debug("Removed footnotes")
        elif footnote_handling == "insert":
            insert_footnotes(doc)
            logger.debug("Inserted footnote content")

        doc.save(source_file)

        logger.info("Utófeldolgozás sikeres", source_file=str(source_file))

    except Exception as e:
        logger.error("Hiba a DOCX előfeldolgozás közben", error=str(e), exc_info=True)
        raise

    logger.info(
        "Utófeldolgozás kész",
        context={
            "source_file": str(source_file),
        },
    )
    return source_file


def docling_process_document(source_file: str) -> str:
    """
    Feldolgoz egy dokumentumot Markdown formátumba a Docling segítségével.

    Args:
        source_file (str): A bemeneti fájl elérési útja (PDF vagy DOCX).

    Returns:
        str: A generált Markdown fájl elérési útja.

    Raises:
        FileNotFoundError: Ha a fájl nem található.
        ValueError: Ha a fájl nem támogatott formátumú.
    """

    logger = structlog_logger.bind(source_file=source_file)
    start_time = time.time()
    logger.info("Docling feldolgozás indítása", context={"source_file": source_file})

    paths = config.settings.get_paths_for_file(source_file)

    validate_document(source_file)

    converter = DocumentConverter()

    try:
        result = converter.convert(source_file)
    except Exception as e:
        logger.error("Hiba a dokumentum konvertálása során", context={"error": str(e)})
        raise RuntimeError(f"A dokumentum konvertálása sikertelen: {str(e)}")

    if not result.document:
        logger.error(
            "A dokumentum konvertálása sikertelen", context={"source_file": source_file}
        )
        raise RuntimeError("A dokumentum konvertálása sikertelen")

    markdown_content = result.document.export_to_markdown()
    logger.debug(
        "Markdown tartalom generálva", context={"content_length": len(markdown_content)}
    )

    output_md_file = paths["markdown_output"]
    with open(output_md_file, "w", encoding="utf-8") as f:
        f.write(markdown_content)

    logger.info(
        "Markdown fájl mentve",
        context={"output_file": output_md_file, "duration": time.time() - start_time},
    )
    return output_md_file


def unstructured_process_markdown(
    source_file: str, markdown_file: str, strategy: str = "auto"
) -> None:
    """
    Markdown fájlt JSON formátumba konvertál az Unstructured könyvtár segítségével.

    Args:
        source_file (str): Az eredeti dokumentum elérési útja.
        markdown_file (str): A bemeneti Markdown fájl elérési útja.
        strategy (str, optional): Feldolgozási stratégia. Alapértelmezett: 'auto'.

    Returns:
        None
    """
    logger = structlog_logger.bind(input_file=markdown_file)
    start_time = time.time()
    logger.info(
        "Unstructured feldolgozás indítása", context={"input_file": markdown_file}
    )

    paths = config.settings.get_paths_for_file(source_file)
    output_json_file = paths["json_output"]

    if not os.path.exists(markdown_file):
        logger.error(
            "Markdown fájl nem található", context={"input_file": markdown_file}
        )
        raise FileNotFoundError(f"A Markdown fájl nem található: {markdown_file}")

    elements = partition(filename=markdown_file, strategy=strategy)
    logger.debug(
        "Dokumentum elemek feldolgozva", context={"element_count": len(elements)}
    )

    write_json([el.to_dict() for el in elements], output_json_file)
    logger.info(
        "JSON fájl mentve",
        context={"output_file": output_json_file, "duration": time.time() - start_time},
    )


def enrich_json(
    source_file: str, input_json_path: str, process_images_with_ai: bool
) -> None:
    """
    JSON tartalom gazdagítása táblázat- és képösszefoglalókkal.

    Args:
        source_file (str): Az eredeti dokumentum elérési útja.
        input_json_path (str): Bemeneti JSON fájl elérési útja.
        process_images_with_ai (bool): Képek AI általi feldolgozása.

    Returns:
        None
    """
    logger = structlog_logger.bind(input_file=input_json_path)
    start_time = time.time()
    logger.info("JSON gazdagítás indítása", context={"input_file": input_json_path})

    paths = config.settings.get_paths_for_file(source_file)
    output_json_path = paths["enriched_json_output"]

    try:
        with open(input_json_path, "r", encoding="utf-8") as f:
            elements = json.load(f)
        logger.debug("JSON fájl betöltve", context={"element_count": len(elements)})
    except FileNotFoundError:
        logger.error("JSON fájl nem található", context={"input_file": input_json_path})
        raise
    except json.JSONDecodeError:
        logger.error(
            "Érvénytelen JSON formátum", context={"input_file": input_json_path}
        )
        raise

    def process_element(element: Dict) -> Dict:
        element_logger = logger.bind(element_type=element["type"])
        try:
            if element["type"] == "Table" and "text_as_html" in element["metadata"]:
                element["metadata"]["table_summary"] = summarize_table(
                    element["metadata"]["text_as_html"]
                )
                element_logger.debug(
                    "Táblázat összefoglalva", context={"element_type": element["type"]}
                )
            elif (
                process_images_with_ai
                and element["type"] == "Image"
                and "image_url" in element["metadata"]
            ):
                element["metadata"]["image_description"] = summarize_image(
                    element["metadata"]["image_url"]
                )
                element_logger.debug(
                    "Kép leírás generálva", context={"element_type": element["type"]}
                )
        except Exception as e:
            element_logger.error(
                "Hiba az elem feldolgozása közben",
                context={"element_type": element["type"], "error": str(e)},
            )
            raise
        return element

    with ThreadPoolExecutor(max_workers=4) as executor:
        elements = list(executor.map(process_element, elements))
    logger.debug("Összes elem feldolgozva", context={"element_count": len(elements)})

    write_json(elements, output_json_path)
    logger.info(
        "Gazdagított JSON mentve",
        context={"output_file": output_json_path, "duration": time.time() - start_time},
    )


def export_text_from_enriched_json(source_file: str, enriched_json_path: str) -> None:
    """
    Szöveg exportálása gazdagított JSON-ból TXT fájlba.

    Args:
        source_file (str): Az eredeti dokumentum elérési útja.
        enriched_json_path (str): Bemeneti gazdagított JSON fájl elérési útja.

    Returns:
        None
    """

    logger = structlog_logger.bind(input_file=enriched_json_path)
    start_time = time.time()
    logger.info(
        "Szöveg exportálása indítása", context={"input_file": enriched_json_path}
    )

    paths = config.settings.get_paths_for_file(source_file)
    output_txt_path = paths["txt_output"]

    try:
        with open(enriched_json_path, "r", encoding="utf-8") as f:
            elements = json.load(f)
        logger.debug("JSON fájl betöltve", context={"element_count": len(elements)})
    except FileNotFoundError:
        logger.error(
            "JSON fájl nem található", context={"input_file": enriched_json_path}
        )
        raise

    paragraphs = []
    for element in elements:
        has_metadata_content = False
        if element.get("metadata"):
            if "table_summary" in element["metadata"]:
                paragraphs.append(element["metadata"]["table_summary"].strip())
                has_metadata_content = True
            if "image_description" in element["metadata"]:
                paragraphs.append(element["metadata"]["image_description"].strip())
                has_metadata_content = True
        if element.get("text") and not has_metadata_content:
            text = element["text"].strip()
            if text:
                paragraphs.append(text)

    logger.debug(
        "Bekezdések összegyűjtve", context={"paragraph_count": len(paragraphs)}
    )

    output_text = "\n\n".join(paragraphs)
    os.makedirs(os.path.dirname(output_txt_path), exist_ok=True)
    with open(output_txt_path, "w", encoding="utf-8") as f:
        f.write(output_text)

    logger.info(
        "Szöveg exportálva",
        context={
            "output_file": output_txt_path,
            "paragraph_count": len(paragraphs),
            "duration": time.time() - start_time,
        },
    )
