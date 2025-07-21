"""
Előfeldolgozási lépések Word dokumentumokhoz.
Tartalmazza a tartalomjegyzék, üres oldalak, fejléc/lábléc eltávolítását és kapcsolódó segédfüggvényeket.
"""

import re
from docx import Document
from docx.oxml.ns import qn


# Előfeldolgozás segédfüggvények


def is_paragraph_empty(paragraph):
    """
    Ellenőrzi, hogy egy bekezdés üres-e (szöveg, run, objektumok, szakaszok alapján).

    Args:
        paragraph (docx.text.Paragraph): A vizsgált bekezdés.

    Returns:
        bool: True, ha a bekezdés üres, különben False.
    """

    if paragraph.text.strip():
        return False

    for run in paragraph.runs:
        if run.text.strip():
            return False

    if len(paragraph._element.xpath(".//w:drawing|.//w:pict|.//w:object")) > 0:
        return False

    if paragraph._element.xpath(".//w:sectPr"):
        return False

    return True


def remove_empty_pages(doc: Document, logger, remove_empty: bool = True) -> int:
    """
    Üres oldalak eltávolítása az oldal tájolásának és elrendezésének megőrzésével.

    Args:
        doc (Document): A Word dokumentum objektum.
        logger: Logger objektum.
        remove_empty (bool): Ha True, eltávolítja az üres oldalakat.

    Returns:
        int: Az eltávolított üres oldalak száma.
    """
    empty_removed = 0

    if not remove_empty:
        return empty_removed

    empty_paragraphs = [para for para in doc.paragraphs if is_paragraph_empty(para)]

    for para in empty_paragraphs:
        parent = para._element.getparent()

        if parent.tag.endswith("sectPr"):
            continue

        try:
            parent.remove(para._element)
            empty_removed += 1
        except Exception as e:
            logger.warning(f"Could not remove paragraph: {str(e)}")
            continue

    return empty_removed


# Tartalomjegyzék kezelés


def remove_toc_by_paragraphs(doc, logger):
    """
    TOC eltávolítása bekezdés stílusok és tartalom alapján.

    Args:
        doc (Document): A Word dokumentum objektum.
        logger: Logger objektum.

    Returns:
        bool: True, ha történt eltávolítás.
    """
    toc_found = False
    paragraphs_to_delete = []

    for para in doc.paragraphs:
        is_toc = (
            para.style.name.lower().startswith(("toc", "tartalomjegyzk", "tj"))
            or para.style.name in ["TJ1", "TJ2", "TJ3", "TJ4", "TJ5", "Jegyzkhivatkozs"]
            or "tartalomjegyzék" in para.text.lower()
            or "Table of Contents" in para.text.lower()
            or any("TOC" in run.text.lower() for run in para.runs)
            or any("PAGEREF" in run.text for run in para.runs)
            or any(
                run._element.find(qn("w:hyperlink")) is not None for run in para.runs
            )
            or any(
                elem.tag == qn("w:fldChar")
                for run in para.runs
                for elem in run._element
            )
            or any("......" in run.text or "." * 10 in run.text for run in para.runs)
            or (
                any(run.text.isdigit() for run in para.runs)
                and any(run._element.find(qn("w:tab")) is not None for run in para.runs)
            )
        )

        if (
            para.style.name == "Listaszerbekezds"
            and para._element.find(qn("w:numPr")) is not None
            and toc_found
        ):
            is_toc = True

        if is_toc:
            toc_found = True
            paragraphs_to_delete.append(para)
        elif toc_found:
            if (
                para.style.name.startswith("Heading")
                and para.text.strip()
                or not para.style.name.lower().startswith(
                    ("tj", "toc", "tartalomjegyzk")
                )
                and para.text.strip()
                and para.style.name != "Listaszerbekezds"
            ):
                break

    for para in paragraphs_to_delete:
        try:
            p = para._element
            if p.getparent() is not None:
                p.getparent().remove(p)
        except Exception as e:
            logger.warning(f"Hiba a bekezdés törlésekor: {str(e)}")

    logger.debug(
        f"Bekezdés alapú TOC eltávolítás: {len(paragraphs_to_delete)} bekezdés törölve"
    )
    return len(paragraphs_to_delete) > 0


def remove_toc_by_field(doc, logger):
    """
    TOC eltávolítása mezőkódok alapján.

    Args:
        doc (Document): A Word dokumentum objektum.
        logger: Logger objektum.

    Returns:
        bool: True, ha történt eltávolítás.
    """
    toc_removed = False
    paragraphs_to_delete = []

    for para_index, para in enumerate(doc.paragraphs):
        for run in para.runs:
            try:
                for elem in run._element:
                    if (
                        (elem.tag == qn("w:fldChar") and "TOC" in elem.getparent().xml)
                        or (
                            elem.tag == qn("w:instrText")
                            and elem.text
                            and ("TOC" in elem.text or "PAGEREF" in elem.text)
                        )
                        or (elem.getparent().tag == qn("w:sdt"))
                    ):
                        paragraphs_to_delete.append(para)
                        logger.debug(f"Mezőkód talált: {para.text[:50]}")
                        toc_removed = True
                        break
            except Exception as elem_error:
                logger.warning(
                    f"Hiba az elemek keresése során a {para_index}. bekezdésben: {str(elem_error)}"
                )

        # SDT elemek keresése
        try:
            sdt_tag = qn("w:sdt")
            sdt = para._element.find(sdt_tag)
            if sdt is not None:
                nsmap = {
                    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                }
                sdt_pr_tag = qn("w:sdtPr")
                sdt_pr = sdt.find(sdt_pr_tag, namespaces=nsmap)
                if sdt_pr is not None:
                    doc_part_obj = sdt_pr.find(qn("w:docPartObj"), namespaces=nsmap)
                    if doc_part_obj is not None:
                        doc_part = doc_part_obj.find(
                            qn("w:docPartGallery"), namespaces=nsmap
                        )
                        if (
                            doc_part is not None
                            and doc_part.get("val") == "Table of Contents"
                        ):
                            paragraphs_to_delete.append(para)
                            logger.debug("SDT blokk megtalálva")
                            toc_removed = True
        except Exception as sdt_error:
            logger.warning(
                f"Hiba az SDT keresése során a {para_index}. bekezdésben: {str(sdt_error)}"
            )

    for para_index, para in enumerate(paragraphs_to_delete):
        try:
            if para._element.getparent() is not None:
                para._element.getparent().remove(para._element)
                logger.debug(f"Bekezdés törölve: {para.text[:50]}")
            else:
                logger.debug(f"Bekezdés már törölve vagy nincs szülő: {para.text[:50]}")
        except Exception as e:
            logger.warning(f"Hiba a bekezdés törlésekor: {str(e)}")

    logger.debug(
        f"Mezőkód alapú TOC eltávolítás: {'sikeres' if toc_removed else 'nem talált TOC-t'}"
    )
    return toc_removed


def remove_toc_by_xml(doc, logger):
    """
    Eltávolítja a tartalomjegyzéket (TOC) a Word dokumentumból az XML struktúra alapján.

    Args:
        doc (Document): python-docx Document objektum
        logger: Logger objektum a naplózáshoz

    Returns:
        bool: True, ha történt eltávolítás.
    """
    # A dokumentum XML gyökere
    document = doc.element.body

    # Lista a törlendő bekezdések tárolására
    paragraphs_to_delete = []

    # Minták a TOC és PAGEREF utasításokhoz
    toc_pattern = re.compile(r"TOC.*")
    pageref_pattern = re.compile(r"PAGEREF _Toc\d+.*")
    toc_hyperlink_pattern = re.compile(r"_TOC_\d+")

    toc_found = False
    paragraph_index = 0

    # Végigiterálunk a dokumentum összes bekezdésén
    for paragraph in document.xpath(".//w:p"):
        is_toc = False
        paragraph_index += 1

        # Ellenőrizzük, hogy a bekezdés tartalmaz-e instrText elemet
        instr_texts = paragraph.xpath(".//w:instrText")
        for instr in instr_texts:
            if instr.text and (
                toc_pattern.match(instr.text) or pageref_pattern.match(instr.text)
            ):
                is_toc = True
                paragraphs_to_delete.append(paragraph)
                logger.info(f"Törlésre jelölt instrText: {instr.text}")
                break

        # Hiperhivatkozások ellenőrzése
        hyperlinks = paragraph.xpath(".//w:hyperlink")
        for hyperlink in hyperlinks:
            anchor = hyperlink.get(qn("w:anchor"))
            if anchor and toc_hyperlink_pattern.match(anchor):
                is_toc = True
                paragraphs_to_delete.append(paragraph)
                logger.info(f"TOC hiperhivatkozás találat: {anchor}")
                break

        # Listaszerbekezds stílus ellenőrzése TOC kontextusban
        style = paragraph.xpath(".//w:pStyle/@w:val")
        if style and style[0] == "Listaszerbekezds" and toc_found:
            is_toc = True
            paragraphs_to_delete.append(paragraph)
            logger.info(f"TOC kontextusban lévő Listaszerbekezds bekezdés: {style[0]}")
        elif style and any(s.startswith(("TJ", "TOC", "toc")) for s in style):
            is_toc = True
            paragraphs_to_delete.append(paragraph)
            logger.info(f"TOC stílusú bekezdés találat: {style[0]}")

        # TOC kontextus frissítése
        if is_toc and not toc_found:
            toc_found = True

        elif toc_found:
            # TOC vége, ha nem TOC stílusú vagy Heading stílusú bekezdést találunk
            if style and (
                style[0].startswith("Heading")
                or (
                    not style[0].lower().startswith(("tj", "toc", "tartalomjegyzk"))
                    and style[0] != "Listaszerbekezds"
                )
            ):
                text_elements = paragraph.xpath(".//w:t")
                if text_elements and any(t.text.strip() for t in text_elements):
                    toc_found = False
                    break

    # SDT elemek kezelése
    sdt_elements = document.xpath(".//w:sdt")
    for sdt in sdt_elements:
        sdt_pr = sdt.find(qn("w:sdtPr"))
        if sdt_pr is not None:
            doc_part_obj = sdt_pr.find(qn("w:docPartObj"))
            if doc_part_obj is not None:
                tag_elem = sdt_pr.find(qn("w:tag"))
                if (
                    tag_elem is not None
                    and "toc" in tag_elem.get(qn("w:val"), "").lower()
                ):
                    parent = sdt.getparent()
                    if parent is not None:
                        parent.remove(sdt)
                        logger.info("TOC SDT elem eltávolítva (tag alapján)")
                        continue

                doc_part = doc_part_obj.find(qn("w:docPartGallery"))
                if (
                    doc_part is not None
                    and doc_part.get(qn("w:val")) == "Table of Contents"
                ):
                    parent = sdt.getparent()
                    if parent is not None:
                        parent.remove(sdt)
                        logger.info("TOC SDT elem eltávolítva")
                        continue

                    for para in sdt.xpath(".//w:p"):
                        paragraphs_to_delete.append(para)

    # Töröljük a megjelölt bekezdéseket
    for paragraph in paragraphs_to_delete:
        try:
            parent = paragraph.getparent()
            if parent is not None:
                parent.remove(paragraph)
                logger.info(
                    "Bekezdés törölve, amely TOC vagy PAGEREF utasítást tartalmazott."
                )
            else:
                logger.warning("Bekezdés szülője nem található, törlés kihagyva.")
        except Exception as e:
            logger.error(f"Hiba a bekezdés törlése közben: {e}")

    # Üres bekezdések eltávolítása
    for paragraph in document.xpath(".//w:p"):
        if not paragraph.xpath(".//w:t") and not paragraph.xpath(".//w:instrText"):
            try:
                parent = paragraph.getparent()
                if parent is not None:
                    parent.remove(paragraph)
            except Exception as e:
                logger.error(f"Hiba az üres bekezdés törlésekor: {e}")

    logger.info("Tartalomjegyzék eltávolítása befejeződött (XML alapú módszer)")
    return len(paragraphs_to_delete) > 0


def remove_toc_by_text(doc, logger, keyword="tartalomjegyzék"):
    """
    TOC eltávolítása szöveges kereséssel.

    Args:
        doc (Document): A Word dokumentum objektum.
        logger: Logger objektum.
        keyword (str): Kulcsszó a kereséshez.

    Returns:
        bool: True, ha történt eltávolítás.
    """
    paragraphs_to_delete = []
    toc_found = False

    for para in doc.paragraphs:
        if keyword in para.text.lower() or "Table of Contents" in para.text.lower():
            toc_found = True
            paragraphs_to_delete.append(para)
        elif toc_found:
            if (
                para.style.name
                in ["TJ1", "TJ2", "TJ3", "TJ4", "TJ5", "Jegyzkhivatkozs", "TOC"]
                or para.style.name.lower().startswith(("toc", "tartalomjegyzk", "TJ"))
                or any("PAGEREF" in run.text for run in para.runs)
                or any(
                    run._element.find(qn("w:hyperlink")) is not None
                    for run in para.runs
                )
                or any(
                    "......" in run.text or "." * 10 in run.text for run in para.runs
                )
            ):
                paragraphs_to_delete.append(para)
            elif (
                para.style.name.startswith("Heading")
                and para.text.strip()
                or para.style.name
                not in ["TJ1", "TJ2", "TJ3", "TJ4", "TJ5", "Jegyzkhivatkozs", "TOC"]
                or para.style.name.lower().startswith(("toc", "tartalomjegyzk", "TJ"))
                and para.text.strip()
            ):
                break

    for para in paragraphs_to_delete:
        p = para._element
        p.getparent().remove(p)

    logger.debug(
        f"Szöveges TOC eltávolítás: {len(paragraphs_to_delete)} bekezdés törölve"
    )
    return len(paragraphs_to_delete) > 0


def remove_toc_by_table(doc, logger):
    """
    TOC eltávolítása táblázatok alapján.

    Args:
        doc (Document): A Word dokumentum objektum.
        logger: Logger objektum.

    Returns:
        bool: True, ha történt eltávolítás.
    """
    toc_removed = False
    for table in doc.tables:
        if any(
            "......" in cell.text or "." * 10 in cell.text
            for row in table.rows
            for cell in row.cells
        ):
            table._element.getparent().remove(table._element)
            toc_removed = True
    return toc_removed
