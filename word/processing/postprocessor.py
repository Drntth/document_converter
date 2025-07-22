"""
Utófeldolgozási lépések Word dokumentumokhoz.
Rövidítések, lábjegyzetek, beszúrása, eltávolítása.
"""

import re
import xml.etree.ElementTree as ET
from typing import Dict, List, Tuple

from config.logging_config import structlog_logger
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt, RGBColor

# Rövidítések kezelése


def extract_abbreviations(doc: Document) -> Dict[str, List[str]]:
    """
    Kinyeri a rövidítéseket és definícióikat a dokumentumból.

    Args:
        doc (Document): Word dokumentum objektum.

    Returns:
        Dict[str, List[str]]: Rövidítés-definíciók szótára, ahol az érték egy definíciók listája.
    """
    abbreviation_dict: Dict[str, List[str]] = {}
    pattern = re.compile(
        r"(.+?)\s*\((?:(?:a\s*)?továbbiakban(?:\s*együtt)?\s*:\s*|továbbiakban\s+)(.+?)\)",
        flags=re.IGNORECASE,
    )

    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        matches = pattern.findall(para.text)
        for full_text, abbrs in matches:
            process_abbreviations(full_text.strip(), abbrs.strip(), abbreviation_dict)

    return abbreviation_dict


def split_abbreviations(abbrs: str) -> List[str]:
    """
    Feldarabolja a rövidítéseket vessző alapján és tisztítja őket.

    Args:
        abbrs (str): A rövidítések szövege (vesszővel elválasztva).

    Returns:
        List[str]: Tisztított rövidítések listája.
    """
    if not abbrs:
        return []
    return [
        abbr.strip().rstrip(".")
        for abbr in re.split(r",\s*|\s*vagy\s*|/\s*", abbrs)
        if abbr.strip()
    ]


def process_abbreviations(
    full_text: str, abbrs: str, abbreviation_dict: Dict[str, List[str]]
) -> None:
    """
    Feldolgozza a rövidítéseket, figyelembe véve a többes rövidítéseket és több definíciót.

    Args:
        full_text (str): A definíció teljes szövege.
        abbrs (str): A rövidítések (vesszővel elválasztva).
        abbreviation_dict (Dict[str, List[str]]): A rövidítés-definíciók szótár.

    Returns:
        None
    """
    abbr_list = split_abbreviations(abbrs)

    for abbr in abbr_list:
        if not abbr:
            continue
        is_valid, definition = validate_abbreviation(abbr, full_text)
        if is_valid:
            # Ha a rövidítés már létezik, hozzáadjuk az új definíciót a listához
            if abbr not in abbreviation_dict:
                abbreviation_dict[abbr] = []
            if definition not in abbreviation_dict[abbr]:
                abbreviation_dict[abbr].append(definition)


def validate_abbreviation(abbr: str, full_text: str) -> Tuple[bool, str]:
    """
    Ellenőrzi, hogy a rövidítés megfelel-e a definíciónak.

    Args:
        abbr (str): A rövidítés.
        full_text (str): A definíció teljes szövege.

    Returns:
        Tuple[bool, str]: (érvényesség, végleges definíció) tuple.
    """
    definition = normalize_text(full_text)
    words = definition.split()

    if is_uppercase_abbr(abbr):
        return validate_uppercase_abbr(abbr, words)
    else:
        return validate_lowercase_abbr(abbr, words)


def is_uppercase_abbr(abbr: str) -> bool:
    """
    Ellenőrzi, hogy a rövidítés nagybetűs-e vagy rövid nagybetűs forma (pl. MNB, ABCDE).

    Args:
        abbr (str): A rövidítés.

    Returns:
        bool: Igaz, ha nagybetűs vagy rövid nagybetűs forma.
    """
    if not abbr:  # Ellenőrizzük, hogy a rövidítés nem üres
        return False
    return abbr.isupper() or (abbr[0].isupper() and len(abbr) <= 5 and abbr.isalpha())


def is_law_or_decree(text: str) -> bool:
    """
    Ellenőrzi, hogy a szöveg törvényre vagy rendeletre utal-e.

    Args:
        text (str): A vizsgálandó szöveg.

    Returns:
        bool: Igaz, ha törvényre vagy rendeletre utal.
    """
    patterns = [
        r"\d{1,4}/\d{4}\.\s*\([IVXLCDM]+\.\s*\d+\.\)",  # Pl. 87/2015. (IV. 9.)
        r"\d{4}\.\s*évi\s+[IVXLCDM]+\.",  # Pl. 2011. évi CCIV.
        r"Korm\.\s*rendelet",  # Korm. rendelet
        r"törvény",  # Törvény
    ]
    return any(re.search(pattern, text, flags=re.IGNORECASE) for pattern in patterns)


def validate_uppercase_abbr(abbr: str, words: List[str]) -> Tuple[bool, str]:
    """
    Validálja a nagybetűs rövidítéseket az iniciálék alapján.

    Args:
        abbr (str): A rövidítés.
        words (List[str]): A definíció szavai.

    Returns:
        Tuple[bool, str]: (érvényesség, végleges definíció) tuple.
    """
    if not abbr:
        return False, ""

    abbr = abbr.upper()
    full_definition = " ".join(words)

    # Szűrjük ki az irreleváns szavakat az iniciálék számításakor
    filtered_words = [w for w in words if w.lower() not in ["és", "a", "az"]]
    filtered_indices = [
        i for i, w in enumerate(words) if w.lower() not in ["és", "a", "az"]
    ]

    # Külön kezelés, ha rendelet vagy törvény
    if is_law_or_decree(full_definition):
        # Ha nincs iniciálé-illeszkedés, használjuk az egész definíciót
        return True, full_definition

    # Legjobb illeszkedés tárolása
    best_match_len = 0
    best_definition = None
    best_sub_words_len = float("inf")

    # Keresés a definícióban olyan kezdőpont után, ahol az iniciálék illeszkednek
    for i in range(len(filtered_words)):
        sub_words = filtered_words[i:]
        sub_initials = [
            get_hungarian_initial(w) for w in sub_words if get_hungarian_initial(w)
        ]
        sub_initial_str = "".join(sub_initials)

        # Pontos illeszkedés ellenőrzése
        if len(sub_initial_str) >= len(abbr) and sub_initial_str.startswith(abbr):
            start_index = filtered_indices[i]
            definition = " ".join(words[start_index:])
            return True, definition

        # Dinamikus részleges illeszkedés az abbr hosszától csökkenően
        for match_len in range(len(abbr), 0, -1):  # Pl. NFTV esetén: 4, 3, 2, 1
            if len(sub_initials) >= match_len:
                # Ellenőrizzük az összes lehetséges kezdőpozíciót az al-szekvenciában
                for start in range(len(sub_initials) - match_len + 1):
                    if all(
                        sub_initials[start + j] == abbr[j] for j in range(match_len)
                    ):
                        start_index = filtered_indices[i]
                        max_words = min(5, len(words[start_index:]))
                        definition = " ".join(
                            words[start_index : start_index + max_words]
                        )
                        # Csak akkor frissítjük, ha jobb az illeszkedés vagy azonos illeszkedésnél rövidebb a definíció
                        if match_len > best_match_len or (
                            match_len == best_match_len
                            and len(sub_words) < best_sub_words_len
                        ):
                            best_match_len = match_len
                            best_definition = definition
                            best_sub_words_len = len(sub_words)

    # Ha találtunk részleges illeszkedést, és az legalább 50%-os
    if best_definition and best_match_len >= 1:
        return True, best_definition

    # Hossz alapú validáció (engedékeny)
    abbr_len = get_adjusted_abbr_len(abbr)
    word_count = len(filtered_words)
    max_diff = min(6, len(abbr) * 2)

    if abs(abbr_len - word_count) <= max_diff:
        max_words = min(5, len(words))
        definition = " ".join(words[-max_words:])
        return True, definition

    return False, " ".join(words)


def validate_lowercase_abbr(abbr: str, words: List[str]) -> Tuple[bool, str]:
    """
    Validálja a kisbetűs rövidítéseket a definíció alapján.

    Args:
        abbr (str): A rövidítés.
        words (List[str]): A definíció szavai.

    Returns:
        Tuple[bool, str]: (érvényesség, végleges definíció) tuple.
    """
    if not abbr:  # Ellenőrizzük, hogy a rövidítés nem üres
        return False, ""

    full_definition = " ".join(words)

    # Normalizáljuk a rövidítést
    normalized_abbr = abbr.lower()
    # Szűrjük ki az irreleváns szavakat az illeszkedés ellenőrzéséhez
    filtered_words = [w for w in words if w.lower() not in ["és", "a", "az"]]
    # Indexek tárolása az eredeti szavakhoz való megfeleltetéshez
    filtered_indices = [
        i for i, w in enumerate(words) if w.lower() not in ["és", "a", "az"]
    ]

    # Külön kezelés, ha rendelet vagy törvény
    if is_law_or_decree(full_definition):
        # Ha nincs iniciálé-illeszkedés, használjuk az egész definíciót
        return True, full_definition

    # Ellenőrizzük, hogy a rövidítés megtalálható-e a szűrt szavak között
    abbr_words = normalized_abbr.split()
    max_words = len(abbr_words) * 4 if len(abbr_words) == 1 else len(abbr_words) * 3
    best_definition = None
    best_match_len = 0

    for i in range(len(filtered_words) - len(abbr_words) + 1):
        sub_words = filtered_words[i : i + len(abbr_words)]

        # Ellenőrizzük, hogy a rövidítés szavai megegyeznek-e az al-szekvenciával
        if all(sub_words[j].lower() == abbr_words[j] for j in range(len(abbr_words))):
            # Az eredeti szavakból állítjuk össze a definíciót, a végétől max. max_words szóval
            start_index = max(
                0, filtered_indices[i] - (max_words - len(abbr_words))
            )  # Az illeszkedés előtt is vehetünk fel szavakat
            end_index = min(
                len(words), start_index + max_words
            )  # Max. max_words szó az start_index-től
            definition_words = words[start_index:end_index]
            definition = " ".join(definition_words)
            match_len = len(abbr_words)

            if match_len > best_match_len:
                best_match_len = match_len
                best_definition = definition

    # Ha találtunk illeszkedést
    if best_definition:
        return True, best_definition

    # Ha nem találtunk pontos illeszkedést, próbáljuk meg a teljes definícióval, korlátozva a max_words-ra
    definition = " ".join(words[max(0, len(words) - max_words) :])
    return True, definition


def normalize_text(text: str) -> str:
    """
    Normalizálja a szöveget, eltávolítva a magyar prefixeket és suffixeket.

    Args:
        text (str): A bemeneti szöveg.

    Returns:
        str: A normalizált szöveg.
    """
    prefixes = ["az ", "a ", "Az ", "A "]
    suffixes = [
        "ának",
        "anak",
        "enek",
        "nak",
        "nek",
        "ban",
        "ben",
        "bol",
        "ből",
        "re",
        "val",
        "vel",
        "ert",
        "ig",
        "hoz",
        "hez",
        "höz",
        "ul",
        "ül",
        "ok",
        "ek",
        "k",
    ]

    for prefix in prefixes:
        if text.lower().startswith(prefix):
            text = text[len(prefix) :].strip()
            break

    for suffix in suffixes:
        if text.lower().endswith(suffix):
            text = text[: -len(suffix)].strip()
            break

    return text


def get_hungarian_initial(word: str) -> str:
    """
    Kinyeri a magyar kezdőbetűt, figyelembe véve a digráfokat és trigrafokat.

    Args:
        word (str): A vizsgálandó szó.

    Returns:
        str: A kezdőbetű, digráf vagy trigraf.
    """
    if not word:
        return ""

    word = word.lower()
    hungarian_trigraphs = ["dzs"]
    hungarian_digraphs = ["cs", "dz", "gy", "ly", "ny", "sz", "ty", "zs"]

    for trigraph in hungarian_trigraphs:
        if word.startswith(trigraph):
            return trigraph.upper()

    for digraph in hungarian_digraphs:
        if word.startswith(digraph):
            return digraph.upper()

    return word[0].upper() if word[0].isalpha() else ""


def get_adjusted_abbr_len(abbr: str) -> int:
    """
    Kiszámítja a rövidítés egységeinek számát (digráfok/trigrafok figyelembevételével).

    Args:
        abbr (str): A rövidítés.

    Returns:
        int: Az egységek száma.
    """
    abbr = abbr.lower()
    hungarian_trigraphs = ["dzs"]
    hungarian_digraphs = ["cs", "dz", "gy", "ly", "ny", "sz", "ty", "zs"]

    units = []
    i = 0
    while i < len(abbr):
        if i + 3 <= len(abbr) and abbr[i : i + 3] in hungarian_trigraphs:
            units.append(abbr[i : i + 3])
            i += 3
        elif i + 2 <= len(abbr) and abbr[i : i + 2] in hungarian_digraphs:
            units.append(abbr[i : i + 2])
            i += 2
        else:
            units.append(abbr[i])
            i += 1
    return len(units)


def remove_abbreviation_phrases(
    doc: Document, abbreviations: Dict[str, List[str]]
) -> None:
    """
    Eltávolítja a '(a továbbiakban: XYZ)' vagy '(továbbiakban XYZ)' részekből a 'továbbiakban' szöveget,
    ha az XYZ rövidítés megtalálható a rövidítések szótárában, de a rövidítést zárójelben megtartja.

    Args:
        doc (Document): A Word dokumentum objektuma.
        abbreviations (Dict[str, List[str]]): Korábban kinyert rövidítések szótára, ahol a kulcs a rövidítés, az érték a teljes kifejezések listája.

    Returns:
        None
    """
    logger = structlog_logger.bind(
        function="remove_abbreviation_phrases",
        total_abbreviations=len(abbreviations),
        operation="phrase_removal",
    )

    pattern = re.compile(
        r"(.+?)\s*\((?:(?:a\s*)?továbbiakban(?:\s*együtt)?\s*:\s*|továbbiakban\s+)(.+?)\)",
        flags=re.IGNORECASE,
    )

    for para in doc.paragraphs:
        new_text = para.text
        matches = pattern.finditer(new_text)

        for match in matches:
            full_match = match.group(0)
            full_expression = match.group(1).strip()
            abbreviation = match.group(2).strip()

            # Ellenőrizzük, hogy a rövidítés szerepel-e a szótárban
            if abbreviation in abbreviations or any(
                abbr in abbreviations for abbr in abbreviation.split(", ")
            ):
                # Csere: csak a teljes kifejezés és a rövidítés marad zárójelben
                new_text = re.sub(
                    re.escape(full_match),
                    f"{full_expression} ({abbreviation})",
                    new_text,
                    flags=re.IGNORECASE,
                )
                logger.info(
                    "Rövidítés kifejezés módosítva",
                    original=full_match,
                    new=f"{full_expression} ({abbreviation})",
                )

        para.text = new_text.strip()

    logger.info("Továbbiakban kifejezések módosítása befejeződött")


def insert_abbreviations(doc: Document, abbr_dict: Dict[str, List[str]]) -> None:
    """
    Rövidítések jelentésének beillesztése ott, ahol a rövidítés önállóan szerepel,
    nincs utána zárójeles definíció, és nem része meglévő zárójeles szerkezetnek.

    Args:
        doc (Document): A Word dokumentum objektuma.
        abbr_dict (Dict[str, List[str]]): A rövidítések szótára, ahol a kulcs a rövidítés, az érték a definíciók listája.

    Returns:
        None
    """
    logger = structlog_logger.bind(
        function="insert_abbreviations",
        total_abbreviations=len(abbr_dict),
        operation="abbreviation_insertion",
    )
    inserted_count = 0

    for para in doc.paragraphs:
        para_text = para.text
        modified = False
        new_text = para_text
        replacements = []

        for abbr, definitions in abbr_dict.items():
            # Ha nincs definíció, ugrunk a következőre
            if not definitions:
                continue
            # Egyszerűség kedvéért az első definíciót használjuk
            full = definitions[0]
            # Regex: rövidítés önálló szóként, utána nem következik zárójel
            pattern = re.compile(
                rf"(?<!\w)\b{re.escape(abbr)}\b(?!\w)(?!\s*\([^)]*\))",
                flags=re.UNICODE | re.IGNORECASE,
            )

            # Find all matches
            for match in pattern.finditer(para_text):
                start, end = match.span()
                abbr_text = match.group(0)

                # Ellenőrizzük, hogy a rövidítés nincs-e zárójelben
                before_text = para_text[:start]
                open_paren_count = before_text.count("(") - before_text.count(")")
                if open_paren_count > 0:
                    logger.debug(
                        "Rövidítés kihagyva, mert zárójelben van",
                        abbreviation=abbr,
                        paragraph_text=para_text[:100] + "..."
                        if len(para_text) > 100
                        else para_text,
                    )
                    continue

                # Ellenőrizzük, hogy a rövidítés nem része-e más szónak
                if (start > 0 and para_text[start - 1].isalnum()) or (
                    end < len(para_text) and para_text[end].isalnum()
                ):
                    logger.debug(
                        "Rövidítés kihagyva, mert része más szónak",
                        abbreviation=abbr,
                        paragraph_text=para_text[:100] + "..."
                        if len(para_text) > 100
                        else para_text,
                    )
                    continue

                # Tároljuk a csere pozícióját és új szövegét
                replacements.append((start, end, f"{abbr_text} ({full})"))
                logger.debug(
                    "Rövidítés jelentése beszúrásra jelölve",
                    abbreviation=abbr,
                    definition=full,
                    paragraph_text=para_text[:100] + "..."
                    if len(para_text) > 100
                    else para_text,
                )

        # Végrehajtjuk a cseréket visszafelé, hogy a pozíciók ne tolódjanak el
        if replacements:
            for start, end, replacement in sorted(replacements, reverse=True):
                new_text = new_text[:start] + replacement + new_text[end:]
                modified = True
                inserted_count += 1

        if modified:
            para.text = new_text
            logger.debug(
                "Bekezdés módosítva",
                paragraph_text=new_text[:100] + "..."
                if len(new_text) > 100
                else new_text,
            )

    logger.info("Rövidítések beszúrása befejeződött", total_inserted=inserted_count)


def prepend_abbreviation_section(
    doc: Document, abbr_dict: Dict[str, List[str]]
) -> None:
    """
    Rövidítések felsorolása a dokumentum elején új bekezdésként, ábécérendben, minden definícióval.

    Args:
        doc (Document): A Word dokumentum objektuma.
        abbr_dict (Dict[str, List[str]]): A rövidítések szótára, ahol a kulcs a rövidítés, az érték a definíciók listája.

    Returns:
        None
    """
    logger = structlog_logger.bind(
        function="prepend_abbreviation_section",
        total_abbreviations=len(abbr_dict),
        operation="abbreviation_section_creation",
    )

    if not abbr_dict:
        logger.info("Üres rövidítési szótár, szekció nem került létrehozásra")
        return

    # "Rövidítések" szekció címe
    section_title = doc.paragraphs[0].insert_paragraph_before("Rövidítések")
    try:
        section_title.style = "Heading 2"
        logger.debug("Rövidítések szekció címe létrehozva 'Heading 2' stílussal")
    except KeyError:
        style = doc.styles.add_style("Heading 2", 1)
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(14)
        font.bold = True
        font.color.rgb = RGBColor(0, 0, 0)  # Fekete szín
        style.paragraph_format.space_before = Pt(18)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        section_title.style = style
        logger.debug("Új 'Heading 2' stílus létrehozva és alkalmazva")

    # Normál stílus a rövidítések listájához
    try:
        normal_style = doc.styles["Normal"]
    except KeyError:
        normal_style = doc.styles.add_style("Normal", 1)
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.space_before = Pt(6)
        normal_style.paragraph_format.space_after = Pt(6)

    # Rövidítések ábécérendben
    sorted_abbrs = sorted(abbr_dict.items(), key=lambda x: x[0].lower())

    # Rövidítések beszúrása, minden definícióval
    for abbr, definitions in sorted_abbrs:
        for i, full in enumerate(definitions, 1):
            para_text = f"{abbr}: {full}" if i == 1 else f"{abbr} ({i}.): {full}"
            para = doc.paragraphs[0].insert_paragraph_before(para_text)
            para.style = normal_style
            logger.debug(
                "Rövidítés hozzáadva a szekcióhoz", abbreviation=abbr, definition=full
            )

    # Üres bekezdés beszúrása a szekció után
    doc.paragraphs[0].insert_paragraph_before("")
    logger.debug("Üres bekezdés hozzáadva a szekció után")

    logger.info("Rövidítések szekció létrehozva", section_length=len(abbr_dict))


# Lábjegyzet kezelése


def remove_footnote_references(doc: Document) -> None:
    """
    Eltávolítja a lábjegyzet hivatkozásokat és szabályos minták alapján a zavaró elemeket.

    Args:
        doc (Document): A Word dokumentum objektuma.

    Returns:
        None
    """
    logger = structlog_logger.bind(
        function="remove_footnote_references", operation="footnote_cleanup"
    )

    for para in doc.paragraphs:
        inline_elements = para._element.xpath(".//w:footnoteReference")
        for el in inline_elements:
            el.getparent().remove(el)

    patterns = [
        r"\[\*\]",  # [*]
        r"\^[0-9]+",  # ^1, ^2
    ]

    for para in doc.paragraphs:
        original = para.text
        cleaned = original
        for pattern in patterns:
            cleaned = re.sub(pattern, "", cleaned)
        if cleaned != original:
            logger.debug(
                f"Lábjegyzet-tisztítás:\nEredeti: {original}\nTisztított: {cleaned}"
            )
        para.text = cleaned


def insert_footnotes(doc: Document) -> None:
    """
    Lábjegyzetek tartalmának beillesztése a dokumentumba.

    Args:
        doc (Document): A Word dokumentum objektuma.

    Returns:
        None
    """

    logger = structlog_logger.bind(
        function="insert_footnotes", operation="footnote_insertion"
    )

    footnote_part = None

    for rel in doc.part.rels.values():
        if rel.reltype == RT.FOOTNOTES:
            footnote_part = rel._target
            break

    if not footnote_part:
        logger.warning("Nem találhatók lábjegyzetek a dokumentumban.")
        return

    footnotes_tree = ET.fromstring(footnote_part.blob)
    nsmap = footnotes_tree.tag.split("}")[0].strip("{")

    footnote_map = {
        fn.attrib.get(f"{{{nsmap}}}id"): "".join(
            node.text or "" for node in fn.iter() if node.text
        )
        for fn in footnotes_tree.findall(".//w:footnote", namespaces={"w": nsmap})
    }

    for para in doc.paragraphs:
        for run in para.runs:
            r_el = run._element
            footnote_refs = r_el.xpath('.//*[local-name()="footnoteReference"]')
            for ref in footnote_refs:
                fid = ref.attrib.get(f"{{{nsmap}}}id")
                footnote_text = footnote_map.get(fid)
                if footnote_text:
                    placeholder = f"<<lábjegyzet: {footnote_text.strip()}>>"
                    run.text += f" {placeholder}"
                    logger.debug(f"Lábjegyzet beszúrva: {placeholder}")
                parent = ref.getparent()
                if parent is not None:
                    parent.remove(ref)
