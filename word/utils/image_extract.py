from pathlib import Path
from docx import Document
from config.logging_config import structlog_logger
import re
from docx.oxml.ns import qn


def extract_images_from_docx(source_file: Path, output_dir: Path) -> int:
    """
    Képek kimentése DOCX-ből az output_dir mappába, placeholder beszúrásával (XML szinten is).
    Args:
        source_file (Path): A bemeneti DOCX fájl.
        output_dir (Path): A képek célmappája.
    Returns:
        int: Kimentett képek száma.
    """
    logger = structlog_logger.bind(source_file=str(source_file))
    doc = Document(source_file)
    output_dir.mkdir(parents=True, exist_ok=True)
    rels = doc.part.rels
    image_map = {}
    parent_map = {}
    # 1. Képek kimentése és parent run eltárolása
    for para in doc.paragraphs:
        for run in para.runs:
            if "graphic" in run._element.xml:
                drawing_elems = run._element.findall(
                    ".//w:drawing", namespaces=run._element.nsmap
                )
                for drawing in drawing_elems:
                    blip = drawing.find(
                        ".//a:blip",
                        namespaces={
                            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
                            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                        },
                    )
                    if blip is not None:
                        rId = blip.get(qn("r:embed"))
                        if rId:
                            parent_map[rId] = run
    image_count = 0
    for rel_id, rel in rels.items():
        if (
            rel.reltype
            == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
        ):
            image_count += 1
            img_data = rel.target_part.blob
            img_ext = rel.target_part.content_type.split("/")[-1]
            img_name = f"image_{image_count}.{img_ext}"
            img_path = output_dir / img_name
            with open(img_path, "wb") as f:
                f.write(img_data)
            image_map[rel_id] = img_name
            # Placeholder beszúrása a parent run után
            if rel_id in parent_map:
                parent_run = parent_map[rel_id]
                para = parent_run._parent
                para.add_run(f" [IMAGE: {img_name}]")
                logger.info(
                    f"Placeholder beszúrva: [IMAGE: {img_name}]",
                    paragraph=str(para.text),
                    output_file=str(source_file),
                )
    logger.info(f"Kimentett képek száma: {image_count}", output_dir=str(output_dir))
    doc.save(source_file)
    return image_count


def replace_image_placeholders_with_markdown(md_file: Path, docname: str) -> None:
    """
    A Markdown fájlban a [IMAGE: ...] placeholdert ![IMAGE](relatív/út) formátumra cseréli.
    Args:
        md_file (Path): A Markdown fájl elérési útja.
        docname (str): A dokumentum neve (kép mappa neve).
    """
    with open(md_file, "r", encoding="utf-8") as f:
        content = f.read()

    def repl(match):
        img_name = match.group(1)
        rel_path = f"images/{docname}/{img_name}"
        return f"![IMAGE]({rel_path})"

    new_content = re.sub(r"\[IMAGE: ([^\]]+)\]", repl, content)

    with open(md_file, "w", encoding="utf-8") as f:
        f.write(new_content)
